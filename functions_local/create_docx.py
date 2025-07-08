import json
from docx import Document

# --- after your chain.predict_and_parse call you have: 
# antrag_obj: PsychotherapieAntrag

# Convert Pydantic model to plain dict
def create_docx(data, doc_path="Psychotherapie_Antrag.docx"):
    # Map field names to their human-readable titles
    titles = {
        'relevante_soziodemographische_daten': 
            'Relevante soziodemographische Daten',
        'symptomatik_psychischer_befund': 
            'Symptomatik und psychischer Befund',
        'somatischer_befund': 
            'Somatischer Befund / Konsiliarbericht',
        'lebensgeschichtliche_angaben':
            'Behandlungsrelevante Angaben zur Lebensgeschichte / Psychosomatik / Systemisches Erklärungsmodell',
        'diagnose_zeitpunkt':
            'Diagnose zum Zeitpunkt der Antragstellung',
        'behandlungsplan_prognose':
            'Behandlungsplan und Prognose',
    }

    # Create the document
    doc = Document()

    # 1) Anrede
    doc.add_paragraph(data['anrede'])

    # 2) Numbered sections 1–6
    for key in titles:
        # Add a numbered heading
        p_heading = doc.add_paragraph(style='List Number')
        p_heading.add_run(f"{titles[key]}")
        # Add the user’s text for that section
        doc.add_paragraph(data[key])

    # 3) Signatur
    doc.add_paragraph(data['signatur'])

    # Save
    doc_path = "Psychotherapie_Antrag.docx"
    doc.save(doc_path)
    print(f"Dokument gespeichert als: {doc_path}")